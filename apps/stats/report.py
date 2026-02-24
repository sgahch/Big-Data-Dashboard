# -*- coding: utf-8 -*-
"""
报告生成模块
支持 Word、PDF、Excel 格式的报告导出
"""
import io
import logging
from datetime import datetime, timedelta
from django.http import HttpResponse
from django.db.models import Count
from django.db.models.functions import TruncMonth
from django.utils import timezone

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

logger = logging.getLogger(__name__)


class ReportGenerator:
    """报告生成器基类"""

    def __init__(self, title=None, subtitle=None):
        self.title = title or '纪检监察数据分析报告'
        self.subtitle = subtitle or ''
        self.generated_at = timezone.now()
        self.data = {}

    def generate(self, **kwargs):
        """生成报告，子类实现"""
        raise NotImplementedError


class WordReportGenerator(ReportGenerator):
    """Word报告生成器"""

    def generate(self, data=None, **kwargs):
        """生成Word报告"""
        doc = Document()

        # 设置中文字体
        self._setup_chinese_fonts(doc)

        # 添加标题
        self._add_title(doc)

        # 添加生成时间
        self._add_meta(doc)

        # 添加报告内容
        if data:
            self._add_summary(doc, data.get('summary', {}))
            self._add_violations_chart(doc, data.get('violations', []))
            self._add_cases_chart(doc, data.get('cases', {}))
            self._add_regions_table(doc, data.get('regions', []))
            self._add_recent_news(doc, data.get('recent_news', []))

        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return output

    def _setup_chinese_fonts(self, doc):
        """设置中文字体"""
        # 为整个文档设置默认字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(12)

    def _add_title(self, doc):
        """添加标题"""
        title = doc.add_heading(self.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.subtitle:
            subtitle = doc.add_paragraph(self.subtitle)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_meta(self, doc):
        """添加元信息"""
        meta = doc.add_paragraph()
        meta.add_run(f'生成时间：{self.generated_at.strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(10)
        meta.add_run('\n')
        meta.add_run(f'数据来源：清风网').font.size = Pt(10)

    def _add_summary(self, doc, summary):
        """添加统计摘要"""
        doc.add_heading('一、数据概览', level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '数值'

        # 数据行
        data = [
            ('总新闻数', str(summary.get('total_news', 0))),
            ('今日新增', str(summary.get('today_news', 0))),
            ('昨日新增', str(summary.get('yesterday_news', 0))),
            ('活跃地区', str(summary.get('active_regions', 0))),
            ('今日爬取', str(summary.get('today_crawled', 0))),
            ('今日新增', str(summary.get('today_new', 0))),
        ]

        for metric, value in data:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = value

    def _add_violations_chart(self, doc, violations):
        """添加违规事项统计"""
        doc.add_heading('二、违规事项分布', level=1)

        if violations:
            # 创建表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '排名'
            hdr_cells[1].text = '违规类型'
            hdr_cells[2].text = '数量'

            for i, item in enumerate(violations[:15], 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = item.get('name', '')
                row[2].text = str(item.get('value', 0))

    def _add_cases_chart(self, doc, cases):
        """添加案件统计"""
        doc.add_heading('三、案件查处趋势', level=1)

        months = cases.get('months', [])
        values = cases.get('values', [])

        if months and values:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '月份'
            hdr_cells[1].text = '案件数'
            hdr_cells[2].text = '环比变化'

            for i, (month, value) in enumerate(zip(months, values)):
                row = table.add_row().cells
                row[0].text = month
                row[1].text = str(value)

                # 计算环比变化
                if i > 0 and values[i-1] > 0:
                    change = ((value - values[i-1]) / values[i-1] * 100)
                    row[2].text = f'{change:+.1f}%'
                else:
                    row[2].text = '-'

    def _add_regions_table(self, doc, regions):
        """添加地区统计表格"""
        doc.add_heading('四、地区分布统计', level=1)

        if regions:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '地区'
            hdr_cells[1].text = '新闻数量'
            hdr_cells[2].text = '占比'

            total = sum(r.get('count', 0) for r in regions)

            for region in regions:
                row = table.add_row().cells
                row[0].text = region.get('region__name', region.get('name', ''))
                row[1].text = str(region.get('count', 0))

                percentage = (region.get('count', 0) / total * 100) if total > 0 else 0
                row[2].text = f'{percentage:.1f}%'

    def _add_recent_news(self, doc, news_list):
        """添加最新新闻列表"""
        doc.add_heading('五、最新通报案例', level=1)

        for i, news in enumerate(news_list[:20], 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. ').bold = True
            p.add_run(news.get('title', ''))

            p.add_run(f'\n    来源：{news.get("region_name", "")} | '
                      f'日期：{news.get("date", "")}').font.size = Pt(10)


class PDFReportGenerator(ReportGenerator):
    """PDF报告生成器"""

    def _get_font_path(self):
        """获取可用的中文字体路径"""
        import os
        import platform

        system = platform.system()

        if system == 'Windows':
            fonts = [
                'C:/Windows/Fonts/msyh.ttc',      # 微软雅黑
                'C:/Windows/Fonts/msyhbd.ttc',    # 微软雅黑粗体
                'C:/Windows/Fonts/simhei.ttf',    # 黑体
                'C:/Windows/Fonts/simsun.ttc',    # 宋体
                'C:/Windows/Fonts/simkai.ttf',    # 楷体
                'C:/Windows/Fonts/arial.ttf',     # Arial (fallback)
            ]
        elif system == 'Darwin':  # macOS
            fonts = [
                '/System/Library/Fonts/PingFang.ttc',
                '/System/Library/Fonts/STHeiti Light.ttc',
            ]
        else:  # Linux
            fonts = [
                '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
                '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttf',
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            ]

        for font_path in fonts:
            normalized = font_path.replace('/', os.sep)
            if os.path.exists(normalized):
                logger.info(f"找到字体文件: {normalized}")
                return normalized
        logger.warning("未找到系统中文字体文件")
        return None

    def generate(self, data=None, **kwargs):
        """生成PDF报告"""
        import os

        # 先生成HTML
        html_content = self._generate_html(data)
        font_path = self._get_font_path()

        # 尝试使用 weasyprint（中文支持更好）
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration

            # 创建字体配置
            font_config = FontConfiguration()
            
            extra_css = ''
            if font_path:
                # 使用系统中可用的字体
                extra_css = f'''
                <style>
                    @font-face {{
                        font-family: "ChineseFont";
                        src: url("{font_path.replace(os.sep, '/')}") format("truetype");
                        font-weight: normal;
                        font-style: normal;
                    }}
                    body, h1, h2, h3, h4, h5, h6, table, th, td {{
                        font-family: "ChineseFont", "Microsoft YaHei", "SimHei", "SimSun", "DejaVu Sans", sans-serif !important;
                    }}
                </style>
                '''
                # 在 HTML 头部插入字体样式
                html_content = html_content.replace('</head>', f'{extra_css}</head>')
            else:
                # 如果没有找到字体，使用默认字体
                html_content = html_content.replace('</head>', '''
                <style>
                    body, h1, h2, h3, h4, h5, h6, table, th, td {
                        font-family: "DejaVu Sans", "Arial Unicode MS", sans-serif !important;
                    }
                </style>
                </head>''')

            output = io.BytesIO()
            HTML(string=html_content, base_url="file://").write_pdf(
                output,
                font_config=font_config
            )
            output.seek(0)
            logger.info("使用 WeasyPrint 生成 PDF")
            return output
        except Exception as weasyprint_error:
            logger.error(f"WeasyPrint failed: {weasyprint_error}")
            # 如果WeasyPrint失败，记录错误并尝试其他方法

        # 尝试使用 xhtml2pdf + ReportLab 注册字体
        try:
            import locale
            import sys
            from xhtml2pdf import pisa
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.lib.fonts import addMapping

            # 设置本地化
            try:
                if sys.platform.startswith('win'):
                    locale.setlocale(locale.LC_ALL, 'Chinese_China.936')  # Windows中文环境
                else:
                    locale.setlocale(locale.LC_ALL, 'zh_CN.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_ALL, 'C.UTF-8')
                except:
                    pass  # 如果设置本地化失败，继续执行

            output = io.BytesIO()

            # 注册中文字体（如果字体文件存在）
            if font_path:
                try:
                    # 尝试注册字体
                    font_name = 'ChineseFont'
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    # 注册字体映射
                    addMapping(font_name, 0, 0, font_name)  # normal
                    addMapping(font_name, 0, 1, font_name)  # italic
                    addMapping(font_name, 1, 0, font_name)  # bold
                    addMapping(font_name, 1, 1, font_name)  # bold italic
                    logger.info(f"成功注册字体: {font_path}")
                    
                    # 在HTML中添加字体样式 - 使用绝对路径（Windows需要file://协议）
                    escaped_font_path = font_path.replace('\\', '/')
                    font_css = f'''
                    <style>
                    @font-face {{ 
                        font-family: 'ChineseFont'; 
                        src: url('file:///{escaped_font_path}'); 
                        font-weight: normal; 
                        font-style: normal; 
                    }}
                    body, h1, h2, h3, h4, h5, h6, table, th, td {{
                        font-family: 'ChineseFont', 'Microsoft YaHei', 'SimHei', 'SimSun', 'DejaVu Sans', sans-serif !important;
                    }}
                    </style>
                    '''
                    html_content = html_content.replace('</head>', f'{font_css}</head>')
                except Exception as font_error:
                    logger.warning(f"字体注册失败: {font_error}")
                    # 如果字体注册失败，使用默认字体样式
                    html_content = html_content.replace('</head>', '''
                    <style>
                    body, h1, h2, h3, h4, h5, h6, table, th, td {
                        font-family: 'Microsoft YaHei', 'SimHei', 'SimSun', 'DejaVu Sans', 'Helvetica', 'Arial', sans-serif !important;
                    }
                    </style>
                    </head>''')
            else:
                # 如果没有找到字体，使用默认字体
                html_content = html_content.replace('</head>', '''
                <style>
                body, h1, h2, h3, h4, h5, h6, table, th, td {
                    font-family: 'Microsoft YaHei', 'SimHei', 'SimSun', 'DejaVu Sans', 'Helvetica', 'Arial', sans-serif !important;
                }
                </style>
                </head>''')

            # 创建PDF
            pdf_status = pisa.CreatePDF(
                html_content,
                dest=output
            )

            # 检查是否有错误
            if pdf_status.err:
                logger.error(f"PDF生成错误: {pdf_status.err}")
                # 即使有错误也尝试返回，因为有时虽然有警告但PDF仍可生成
                if output.tell() > 0:
                    output.seek(0)
                    logger.info("使用 xhtml2pdf 生成 PDF，尽管有警告")
                    return output
                else:
                    raise Exception(f"PDF生成错误码: {pdf_status.err}")

            output.seek(0)
            logger.info("使用 xhtml2pdf 生成 PDF")
            return output

        except ImportError as e:
            logger.error(f"缺少依赖: {e}")
            raise ImportError("请安装依赖: pip install weasyprint xhtml2pdf reportlab")
        except Exception as e:
            logger.error(f"PDF生成失败: {e}")
            import traceback
            traceback.print_exc()
            raise

    def _generate_html(self, data):
        """生成报告HTML"""
        summary = data.get('summary', {})
        violations = data.get('violations', [])
        cases = data.get('cases', {})
        regions = data.get('regions', [])
        recent_news = data.get('recent_news', [])

        # 构建表格行
        violations_rows = ''.join(
            f"<tr><td>{i+1}</td><td>{v.get('name', '')}</td><td>{v.get('value', 0)}</td></tr>"
            for i, v in enumerate(violations[:15])
        )

        cases_rows = ''.join(
            f"<tr><td>{m}</td><td>{v}</td><td>{self._calc_change(cases.get('values', []), i)}</td></tr>"
            for i, (m, v) in enumerate(zip(cases.get('months', []), cases.get('values', [])))
        )

        # 计算地区占比
        total_region_count = sum(r.get('count', 0) for r in regions)
        regions_rows = ''.join(
            f"<tr><td>{r.get('region__name', r.get('name', ''))}</td>"
            f"<td>{r.get('count', 0)}</td>"
            f"<td>{self._calc_percentage(total_region_count, r.get('count', 0))}%</td></tr>"
            for r in regions
        )

        news_items = ''.join(
            f"<div class='news-item'><strong>{n.get('title', '')}</strong>"
            f"<br><small>来源：{n.get('region_name', '')} | 日期：{n.get('date', '')}</small></div>"
            for n in recent_news[:20]
        )

        html = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <title>{self.title}</title>
            <style>
                @page {{ size: A4; margin: 2cm }}
                body {{
                    font-family: 'Microsoft YaHei', 'SimSun', '宋体', sans-serif;
                    font-size: 12px;
                    color: #333;
                    line-height: 1.6
                }}
                h1 {{ color: #1a4a8c; text-align: center; font-size: 24px; margin-bottom: 10px }}
                h2 {{ color: #2c5aa0; border-bottom: 2px solid #2c5aa0; padding-bottom: 8px; margin-top: 30px }}
                h3 {{ color: #3d6cb8; margin-top: 20px }}
                .meta {{ text-align: center; color: #666; font-size: 10px; margin-bottom: 30px }}
                table {{ width: 100%; border-collapse: collapse; margin: 15px 0 }}
                th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left }}
                th {{ background-color: #00D2FF; color: #fff; font-weight: bold }}
                tr:nth-child(even) {{ background-color: #f9f9f9 }}
                .summary-grid {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 15px; margin: 20px 0 }}
                .summary-card {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; text-align: center }}
                .summary-card .value {{ font-size: 28px; font-weight: bold }}
                .summary-card .label {{ font-size: 12px; opacity: 0.9 }}
                .news-item {{ padding: 10px 0; border-bottom: 1px solid #eee }}
                .news-item:last-child {{ border-bottom: none }}
            </style>
        </head>
        <body>
            <h1>{self.title}</h1>
            <div class="meta">
                <p>生成时间：{self.generated_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
                <p>数据来源：清风网</p>
            </div>

            <h2>一、数据概览</h2>
            <div class="summary-grid">
                <div class="summary-card">
                    <div class="value">{summary.get('total_news', 0)}</div>
                    <div class="label">总新闻数</div>
                </div>
                <div class="summary-card">
                    <div class="value">{summary.get('today_news', 0)}</div>
                    <div class="label">今日新增</div>
                </div>
                <div class="summary-card">
                    <div class="value">{summary.get('yesterday_news', 0)}</div>
                    <div class="label">昨日新增</div>
                </div>
                <div class="summary-card">
                    <div class="value">{summary.get('active_regions', 0)}</div>
                    <div class="label">活跃地区</div>
                </div>
                <div class="summary-card">
                    <div class="value">{summary.get('today_crawled', 0)}</div>
                    <div class="label">今日爬取</div>
                </div>
                <div class="summary-card">
                    <div class="value">{summary.get('today_new', 0)}</div>
                    <div class="label">今日新增</div>
                </div>
            </div>

            <h2>二、违规事项分布</h2>
            <table>
                <tr><th>排名</th><th>违规类型</th><th>数量</th></tr>
                {violations_rows}
            </table>

            <h2>三、案件查处趋势</h2>
            <table>
                <tr><th>月份</th><th>案件数</th><th>环比变化</th></tr>
                {cases_rows}
            </table>

            <h2>四、地区分布统计</h2>
            <table>
                <tr><th>地区</th><th>新闻数量</th><th>占比</th></tr>
                {regions_rows}
            </table>

            <h2>五、最新通报案例</h2>
            {news_items}
        </body>
        </html>
        """
        return html

    def _calc_change(self, values, index):
        """计算环比变化"""
        if index > 0 and values[index-1] > 0:
            change = (values[index] - values[index-1]) / values[index-1] * 100
            return f'{change:+.1f}%'
        return '-'

    def _calc_percentage(self, total, count):
        """计算占比"""
        if total > 0:
            return f'{count / total * 100:.1f}'
        return '0'


class ExcelReportGenerator(ReportGenerator):
    """Excel报告生成器"""

    def generate(self, data=None, **kwargs):
        """生成Excel报告"""
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()

        # Sheet 1: 统计概览
        ws1 = wb.active
        ws1.title = '统计概览'
        self._setup_sheet(ws1)
        self._fill_summary(ws1, data.get('summary', {}))

        # Sheet 2: 违规事项分布
        ws2 = wb.create_sheet('违规事项分布')
        self._setup_sheet(ws2)
        self._fill_list(ws2, data.get('violations', []), ['违规类型', '数量'])

        # Sheet 3: 案件趋势
        ws3 = wb.create_sheet('案件趋势')
        self._setup_sheet(ws3)
        self._fill_cases_trend(ws3, data.get('cases', {}))

        # Sheet 4: 地区统计
        ws4 = wb.create_sheet('地区统计')
        self._setup_sheet(ws4)
        self._fill_regions(ws4, data.get('regions', []))

        # Sheet 5: 最新新闻
        ws5 = wb.create_sheet('最新新闻')
        self._setup_sheet_with_extra_columns(ws5)  # 设置更多的列宽
        self._fill_news(ws5, data.get('recent_news', []))

        # Sheet 6: 详细统计
        ws6 = wb.create_sheet('详细统计')
        self._setup_sheet(ws6)
        self._fill_detailed_stats(ws6, data)

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return output

    def _fill_cases_trend(self, ws, cases):
        """填充案件趋势数据"""
        ws['A1'] = '月份'
        ws['B1'] = '案件数'
        ws['C1'] = '环比变化'
        ws['D1'] = '同比增长'

        months = cases.get('months', [])
        values = cases.get('values', [])

        for i, (month, value) in enumerate(zip(months, values), 2):
            ws[f'A{i}'] = month
            ws[f'B{i}'] = value
            
            # 环比变化
            if i > 2 and len(values) > i-3 and values[i-3] > 0:
                change = ((value - values[i-3]) / values[i-3] * 100)
                ws[f'C{i}'] = f'{change:+.1f}%'
            else:
                ws[f'C{i}'] = '-'
            
            # 同比变化（与去年同期比较）
            # 简化处理：假设有足够数据
            ws[f'D{i}'] = '-'  # 可根据实际数据计算

    def _setup_sheet_with_extra_columns(self, ws):
        """设置工作表样式（适用于更多列）"""
        from openpyxl.styles import Border, Side  # 在函数内部导入，避免名称冲突
        
        header_font = Font(bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='00D2FF', end_color='00D2FF', fill_type='solid')
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border

        # 设置列宽
        ws.column_dimensions['A'].width = 8   # 序号
        ws.column_dimensions['B'].width = 50  # 标题
        ws.column_dimensions['C'].width = 15  # 地区
        ws.column_dimensions['D'].width = 12  # 日期
        ws.column_dimensions['E'].width = 30  # 标签
        ws.column_dimensions['F'].width = 20  # 来源

    def _fill_detailed_stats(self, ws, data):
        """填充详细统计数据"""
        ws['A1'] = '统计项目'
        ws['B1'] = '数值'
        ws['C1'] = '说明'

        row = 2
        
        # 添加更多统计信息
        summary = data.get('summary', {})
        ws[f'A{row}'] = '总新闻数'
        ws[f'B{row}'] = summary.get('total_news', 0)
        ws[f'C{row}'] = '系统中所有已发布新闻的总数'
        row += 1

        ws[f'A{row}'] = '今日新增'
        ws[f'B{row}'] = summary.get('today_news', 0)
        ws[f'C{row}'] = '今天新增的新闻数量'
        row += 1

        ws[f'A{row}'] = '昨日新增'
        ws[f'B{row}'] = summary.get('yesterday_news', 0)
        ws[f'C{row}'] = '昨天新增的新闻数量'
        row += 1

        ws[f'A{row}'] = '活跃地区数'
        ws[f'B{row}'] = summary.get('active_regions', 0)
        ws[f'C{row}'] = '有新闻发布的地区数量'
        row += 1

        ws[f'A{row}'] = '今日爬取'
        ws[f'B{row}'] = summary.get('today_crawled', 0)
        ws[f'C{row}'] = '今天爬取的新闻数量'
        row += 1

        ws[f'A{row}'] = '今日新发现'
        ws[f'B{row}'] = summary.get('today_new', 0)
        ws[f'C{row}'] = '今天新发现的新闻数量'
        row += 1

        # 添加违规类型统计
        violations = data.get('violations', [])
        if violations:
            ws[f'A{row}'] = '主要违规类型'
            ws[f'B{row}'] = len(violations)
            top_violations = [v.get('name', '') for v in violations[:3]]
            ws[f'C{row}'] = '、'.join(top_violations) if top_violations else '无'
            row += 1

        # 添加地区统计
        regions = data.get('regions', [])
        if regions:
            ws[f'A{row}'] = '主要活跃地区'
            ws[f'B{row}'] = len(regions)
            top_regions = [r.get('region__name', r.get('name', '')) for r in regions[:3]]
            ws[f'C{row}'] = '、'.join(top_regions) if top_regions else '无'
            row += 1

        # 添加数据时间范围
        news_list = data.get('recent_news', [])
        if news_list:
            dates = [n.get('date') for n in news_list if n.get('date')]
            if dates:
                from datetime import datetime
                dates = [d for d in dates if d]
                if dates:
                    # 假设日期格式为字符串，尝试解析
                    try:
                        date_objects = [datetime.fromisoformat(str(d).replace('Z', '+00:00')) if 'T' in str(d) else datetime.strptime(str(d), '%Y-%m-%d') for d in dates]
                        ws[f'A{row}'] = '数据时间范围'
                        ws[f'B{row}'] = f'{min(date_objects).strftime("%Y-%m-%d")} 至 {max(date_objects).strftime("%Y-%m-%d")}'
                        ws[f'C{row}'] = '当前报告覆盖的时间段'
                    except:
                        pass

    def _setup_sheet(self, ws):
        """设置工作表样式"""
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color='00D2FF', end_color='00D2FF', fill_type='solid')

        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')

        # 设置列宽
        ws.column_dimensions['A'].width = 30
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 20

    def _fill_summary(self, ws, summary):
        """填充统计摘要"""
        data = [
            ('总新闻数', summary.get('total_news', 0)),
            ('今日新增', summary.get('today_news', 0)),
            ('昨日新增', summary.get('yesterday_news', 0)),
            ('活跃地区', summary.get('active_regions', 0)),
            ('今日爬取', summary.get('today_crawled', 0)),
            ('今日新增', summary.get('today_new', 0)),
        ]

        for i, (metric, value) in enumerate(data, 1):
            ws[f'A{i}'] = metric
            ws[f'B{i}'] = value

    def _fill_list(self, ws, items, headers):
        """填充列表数据"""
        ws['A1'] = headers[0]
        ws['B1'] = headers[1]

        for i, item in enumerate(items, 2):
            ws[f'A{i}'] = item.get('name', '')
            ws[f'B{i}'] = item.get('value', 0) or item.get('count', 0)

    def _fill_regions(self, ws, regions):
        """填充地区统计"""
        ws['A1'] = '地区'
        ws['B1'] = '数量'
        ws['C1'] = '占比'

        total = sum(r.get('count', 0) for r in regions)

        for i, region in enumerate(regions, 2):
            ws[f'A{i}'] = region.get('region__name', region.get('name', ''))
            ws[f'B{i}'] = region.get('count', 0)
            percentage = (region.get('count', 0) / total * 100) if total > 0 else 0
            ws[f'C{i}'] = f'{percentage:.1f}%'

    def _fill_news(self, ws, news_list):
        """填充新闻列表"""
        ws['A1'] = '序号'
        ws['B1'] = '标题'
        ws['C1'] = '地区'
        ws['D1'] = '日期'
        ws['E1'] = '标签'

        for i, news in enumerate(news_list, 2):
            ws[f'A{i}'] = i - 1
            ws[f'B{i}'] = news.get('title', '')
            ws[f'C{i}'] = news.get('region_name', '')
            ws[f'D{i}'] = str(news.get('date', ''))
            ws[f'E{i}'] = news.get('tag_names', '')


def generate_report(report_type='word', data=None, **kwargs):
    """
    生成报告的统一入口

    Args:
        report_type: 报告类型 ('word', 'excel', 'pdf')
        data: 报告数据
        **kwargs: 其他参数

    Returns:
        文件流
    """
    generators = {
        'word': WordReportGenerator,
        'excel': ExcelReportGenerator,
        'pdf': PDFReportGenerator,
    }

    generator_class = generators.get(report_type, WordReportGenerator)
    generator = generator_class(
        title=kwargs.get('title', '纪检监察数据分析报告'),
        subtitle=kwargs.get('subtitle', '')
    )

    return generator.generate(data)


def get_report_data(start_date=None, end_date=None, region=None):
    """
    获取报告所需数据

    Args:
        start_date: 开始日期
        end_date: 结束日期
        region: 地区筛选

    Returns:
        报告数据字典
    """
    from apps.news.models import News, Region, Tag, CrawlLog
    from apps.crawler.crawler import REGIONS

    today = timezone.now().date()

    # 统计摘要
    total_news = News.objects.filter(status='published').count()
    today_news = News.objects.filter(crawl_time__date=today).count()
    yesterday_news = News.objects.filter(
        crawl_time__date=today - timedelta(days=1)
    ).count()

    today_logs = CrawlLog.objects.filter(crawl_time__date=today)
    today_crawled = today_logs.aggregate(total=Count('total_crawled'))['total'] or 0
    today_new = today_logs.aggregate(total=Count('new_count'))['total'] or 0

    # 违规事项
    violations = {}
    news_qs = News.objects.filter(status='published')
    for n in news_qs:
        tags = n.get_tag_names_list()
        for tag in tags:
            violations[tag] = violations.get(tag, 0) + 1
    violations_list = [
        {'name': k, 'value': v}
        for k, v in sorted(violations.items(), key=lambda x: x[1], reverse=True)
    ]

    # 案件趋势（近12月）
    cases_stats = list((News.objects
                   .filter(status='published')
                   .annotate(month=TruncMonth('crawl_time'))
                   .values('month')
                   .annotate(count=Count('id'))
                   .order_by('month')))[-12:]

    cases = {
        'months': [item['month'].strftime('%Y-%m') for item in cases_stats],
        'values': [item['count'] for item in cases_stats]
    }

    # 地区统计
    regions_stats = (News.objects
                     .filter(status='published')
                     .values('region__name', 'region__code')
                     .annotate(count=Count('id'))
                     .order_by('-count'))

    regions_list = list(regions_stats)

    # 最新新闻
    recent_news = News.objects.filter(
        status='published'
    ).select_related('region').prefetch_related('tags')[:50]

    recent_news_data = [
        {
            'title': n.title,
            'region_name': n.region_name,
            'date': str(n.date) if n.date else '',
            'tag_names': n.tag_names,
            'url': n.url
        }
        for n in recent_news
    ]

    return {
        'summary': {
            'total_news': total_news,
            'today_news': today_news,
            'yesterday_news': yesterday_news,
            'active_regions': regions_stats.values('region').distinct().count(),
            'today_crawled': today_crawled,
            'today_new': today_new,
        },
        'violations': violations_list,
        'cases': cases,
        'regions': regions_list,
        'recent_news': recent_news_data
    }
